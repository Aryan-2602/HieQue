"""
Enhanced document processing and text extraction for HieQue framework
Supports multiple document formats with intelligent text processing
"""

import os
import re
import time
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import logging

import PyPDF2
from docx import Document
import pandas as pd
import numpy as np

from ..logger import get_logger
from ..exception import DocumentProcessingError
from ..utils import chunk_text, normalize_text, extract_metadata_from_filename, generate_document_id

logger = get_logger(__name__)

class DocumentProcessor:
    """
    Multi-format document processor for academic content
    """
    
    def __init__(self, 
                 chunk_size: int = 1000,
                 overlap: int = 200,
                 min_chunk_length: int = 100):
        """
        Initialize document processor
        
        Args:
            chunk_size: Maximum size of text chunks
            overlap: Overlap between chunks
            min_chunk_length: Minimum length for a valid chunk
        """
        self.chunk_size = chunk_size
        self.overlap = overlap
        self.min_chunk_length = min_chunk_length
        self.logger = logger
        
        # Supported file extensions
        self.supported_extensions = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.txt': self._process_txt,
            '.csv': self._process_csv,
            '.xlsx': self._process_excel
        }
    
    def process_document(self, file_path: str, **kwargs) -> List[Dict[str, Any]]:
        """
        Process a single document and return chunks with metadata
        
        Args:
            file_path: Path to the document
            **kwargs: Additional processing options
            
        Returns:
            List of document chunks with metadata
        """
        start_time = time.time()
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise DocumentProcessingError(f"File not found: {file_path}")
        
        file_ext = file_path.suffix.lower()
        if file_ext not in self.supported_extensions:
            raise DocumentProcessingError(f"Unsupported file type: {file_ext}")
        
        self.logger.info(f"Processing document: {file_path.name}")
        
        try:
            # Extract text content
            text_content = self.supported_extensions[file_ext](file_path, **kwargs)
            
            # Extract metadata
            metadata = self._extract_document_metadata(file_path, text_content)
            
            # Chunk the text
            chunks = self._create_chunks(text_content, metadata)
            
            # Generate document IDs and finalize
            documents = self._finalize_chunks(chunks, metadata)
            
            processing_time = time.time() - start_time
            self.logger.info(f"Document processed in {processing_time:.2f}s: {len(documents)} chunks created")
            
            return documents
            
        except Exception as e:
            self.logger.error(f"Failed to process document {file_path}: {e}")
            raise DocumentProcessingError(f"Document processing failed: {e}")
    
    def process_directory(self, directory_path: str, **kwargs) -> List[Dict[str, Any]]:
        """
        Process all supported documents in a directory
        
        Args:
            directory_path: Path to directory
            **kwargs: Additional processing options
            
        Returns:
            List of all document chunks from the directory
        """
        directory_path = Path(directory_path)
        if not directory_path.exists() or not directory_path.is_dir():
            raise DocumentProcessingError(f"Directory not found: {directory_path}")
        
        all_documents = []
        processed_files = 0
        
        for file_path in directory_path.rglob("*"):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_extensions:
                try:
                    documents = self.process_document(str(file_path), **kwargs)
                    all_documents.extend(documents)
                    processed_files += 1
                    self.logger.info(f"Processed {file_path.name}: {len(documents)} chunks")
                except Exception as e:
                    self.logger.warning(f"Failed to process {file_path.name}: {e}")
                    continue
        
        self.logger.info(f"Directory processing completed: {processed_files} files, {len(all_documents)} total chunks")
        return all_documents
    
    def _process_pdf(self, file_path: Path, **kwargs) -> str:
        """Extract text from PDF file"""
        try:
            text_content = ""
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text.strip():
                        # Add page separator
                        text_content += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
            
            if not text_content.strip():
                raise DocumentProcessingError("No text content extracted from PDF")
            
            return text_content
            
        except Exception as e:
            raise DocumentProcessingError(f"PDF processing failed: {e}")
    
    def _process_docx(self, file_path: Path, **kwargs) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text_content = ""
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content += para.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells if cell.text.strip()])
                    if row_text.strip():
                        text_content += row_text + "\n"
            
            if not text_content.strip():
                raise DocumentProcessingError("No text content extracted from DOCX")
            
            return text_content
            
        except Exception as e:
            raise DocumentProcessingError(f"DOCX processing failed: {e}")
    
    def _process_txt(self, file_path: Path, **kwargs) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text_content = file.read()
            
            if not text_content.strip():
                raise DocumentProcessingError("No text content in TXT file")
            
            return text_content
            
        except Exception as e:
            raise DocumentProcessingError(f"TXT processing failed: {e}")
    
    def _process_csv(self, file_path: Path, **kwargs) -> str:
        """Extract text from CSV file"""
        try:
            df = pd.read_csv(file_path)
            text_content = ""
            
            # Add column headers
            text_content += "Columns: " + ", ".join(df.columns.tolist()) + "\n\n"
            
            # Add sample data
            for idx, row in df.head(100).iterrows():  # Limit to first 100 rows
                row_text = " | ".join([str(val) for val in row.values])
                text_content += f"Row {idx + 1}: {row_text}\n"
            
            if len(df) > 100:
                text_content += f"\n... and {len(df) - 100} more rows\n"
            
            return text_content
            
        except Exception as e:
            raise DocumentProcessingError(f"CSV processing failed: {e}")
    
    def _process_excel(self, file_path: Path, **kwargs) -> str:
        """Extract text from Excel file"""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            text_content = ""
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                text_content += f"\n--- Sheet: {sheet_name} ---\n"
                text_content += "Columns: " + ", ".join(df.columns.tolist()) + "\n\n"
                
                # Add sample data
                for idx, row in df.head(50).iterrows():  # Limit to first 50 rows per sheet
                    row_text = " | ".join([str(val) for val in row.values])
                    text_content += f"Row {idx + 1}: {row_text}\n"
                
                if len(df) > 50:
                    text_content += f"\n... and {len(df) - 50} more rows\n"
            
            return text_content
            
        except Exception as e:
            raise DocumentProcessingError(f"Excel processing failed: {e}")
    
    def _extract_document_metadata(self, file_path: Path, text_content: str) -> Dict[str, Any]:
        """Extract comprehensive metadata from document"""
        metadata = {
            'source_file': str(file_path),
            'filename': file_path.name,
            'file_size_mb': file_path.stat().st_size / (1024 * 1024),
            'file_type': file_path.suffix.lower(),
            'processing_timestamp': time.time(),
            'total_characters': len(text_content),
            'total_words': len(text_content.split()),
            'estimated_pages': self._estimate_pages(text_content)
        }
        
        # Extract additional metadata from filename
        filename_metadata = extract_metadata_from_filename(file_path.name)
        metadata.update(filename_metadata)
        
        # Extract content-based metadata
        content_metadata = self._extract_content_metadata(text_content)
        metadata.update(content_metadata)
        
        return metadata
    
    def _estimate_pages(self, text_content: str) -> int:
        """Estimate number of pages based on text content"""
        # Rough estimate: 2500 characters per page
        estimated_pages = max(1, len(text_content) // 2500)
        return estimated_pages
    
    def _extract_content_metadata(self, text_content: str) -> Dict[str, Any]:
        """Extract metadata from content analysis"""
        metadata = {}
        
        # Detect language (simple heuristic)
        metadata['language'] = self._detect_language(text_content)
        
        # Detect academic content indicators
        metadata['academic_indicators'] = self._detect_academic_content(text_content)
        
        # Extract potential topics
        metadata['potential_topics'] = self._extract_potential_topics(text_content)
        
        return metadata
    
    def _detect_language(self, text: str) -> str:
        """Simple language detection"""
        # Count common English words
        english_words = {'the', 'and', 'of', 'to', 'in', 'a', 'is', 'that', 'it', 'with'}
        words = set(text.lower().split())
        english_count = len(words.intersection(english_words))
        
        if english_count > 5:
            return 'english'
        else:
            return 'unknown'
    
    def _detect_academic_content(self, text: str) -> List[str]:
        """Detect indicators of academic content"""
        indicators = []
        
        # Check for academic patterns
        if re.search(r'\b(abstract|introduction|conclusion|references|bibliography)\b', text.lower()):
            indicators.append('academic_structure')
        
        if re.search(r'\b\d{4}\b', text):  # Year references
            indicators.append('year_references')
        
        if re.search(r'\b(et al\.|ibid\.|op\. cit\.)\b', text.lower()):
            indicators.append('citation_style')
        
        if re.search(r'[A-Z][a-z]+ et al\.', text):
            indicators.append('author_citations')
        
        return indicators
    
    def _extract_potential_topics(self, text: str) -> List[str]:
        """Extract potential topics from text"""
        topics = []
        
        # Look for capitalized phrases (potential proper nouns/topics)
        capitalized_phrases = re.findall(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b', text)
        
        # Filter and count
        topic_counts = {}
        for phrase in capitalized_phrases:
            if len(phrase.split()) <= 4:  # Limit to reasonable phrase length
                topic_counts[phrase] = topic_counts.get(phrase, 0) + 1
        
        # Get top topics
        sorted_topics = sorted(topic_counts.items(), key=lambda x: x[1], reverse=True)
        topics = [topic for topic, count in sorted_topics[:10] if count > 1]
        
        return topics
    
    def _create_chunks(self, text_content: str, metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Create text chunks with metadata"""
        # Normalize text
        normalized_text = normalize_text(text_content)
        
        # Create chunks
        text_chunks = chunk_text(
            normalized_text, 
            chunk_size=self.chunk_size, 
            overlap=self.overlap
        )
        
        chunks = []
        for i, chunk in enumerate(text_chunks):
            if len(chunk.strip()) >= self.min_chunk_length:
                chunk_metadata = metadata.copy()
                chunk_metadata.update({
                    'chunk_index': i,
                    'chunk_size': len(chunk),
                    'chunk_start_char': normalized_text.find(chunk),
                    'chunk_end_char': normalized_text.find(chunk) + len(chunk)
                })
                
                chunks.append({
                    'content': chunk,
                    'metadata': chunk_metadata
                })
        
        return chunks
    
    def _finalize_chunks(self, chunks: List[Dict[str, Any]], metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Finalize chunks with document IDs and final metadata"""
        documents = []
        
        for i, chunk in enumerate(chunks):
            # Generate unique document ID
            doc_id = generate_document_id(chunk['content'], chunk['metadata'])
            
            # Create final document structure
            document = {
                'id': doc_id,
                'content': chunk['content'],
                'metadata': chunk['metadata'],
                'chunk_number': i + 1,
                'total_chunks': len(chunks)
            }
            
            documents.append(document)
        
        return documents
    
    def get_processing_stats(self, documents: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Get statistics about processed documents"""
        if not documents:
            return {}
        
        total_chunks = len(documents)
        total_characters = sum(len(doc['content']) for doc in documents)
        total_words = sum(len(doc['content'].split()) for doc in documents)
        
        # File type distribution
        file_types = {}
        for doc in documents:
            file_type = doc['metadata'].get('file_type', 'unknown')
            file_types[file_type] = file_types.get(file_type, 0) + 1
        
        # Chunk size distribution
        chunk_sizes = [len(doc['content']) for doc in documents]
        
        stats = {
            'total_documents': total_chunks,
            'total_characters': total_characters,
            'total_words': total_words,
            'average_chunk_size': np.mean(chunk_sizes) if chunk_sizes else 0,
            'min_chunk_size': min(chunk_sizes) if chunk_sizes else 0,
            'max_chunk_size': max(chunk_sizes) if chunk_sizes else 0,
            'file_type_distribution': file_types,
            'processing_timestamp': time.time()
        }
        
        return stats
